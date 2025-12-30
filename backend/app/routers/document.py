"""文档处理相关API路由"""
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from ..models.schemas import FileUploadResponse, AnalysisRequest, AnalysisType, WordExportRequest
from ..services.file_service import FileService
# from ..services.openai_service import OpenAIService
from ..services.qwen_openai_service import OpenAIService
from ..utils.config_manager import config_manager
from ..utils.sse import sse_response
import json
import io
import re
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from urllib.parse import quote

router = APIRouter(prefix="/api/document", tags=["文档处理"])


def set_run_font_simsun(run: docx.text.run.Run) -> None:
    """统一将 run 字体设置为宋体（包含 EastAsia 字体设置）"""
    run.font.name = "宋体"
    r = run._element.rPr
    if r is not None and r.rFonts is not None:
        r.rFonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_font_simsun(paragraph: docx.text.paragraph.Paragraph) -> None:
    """将段落内所有 runs 字体设置为宋体"""
    for run in paragraph.runs:
        set_run_font_simsun(run)


def clean_html_tags(content: str) -> str:
    """清理HTML标签，保留Markdown格式的内容"""
    import re
    
    if not content:
        return ''
    
    # 首先处理HTML实体，确保所有实体都被正确转换
    html_entities = {
        r'&nbsp;': ' ',
        r'&amp;': '&',
        r'&lt;': '<',
        r'&gt;': '>',
        r'&quot;': '"',
        r'&#39;': "'",
        r'&copy;': '©',
        r'&reg;': '®',
        r'&trade;': '™',
        r'&ldquo;': '"',
        r'&rdquo;': '"',
        r'&lsquo;': "'",
        r'&rsquo;': "'",
        r'&ndash;': '-',
        r'&mdash;': '--',
        r'&plusmn;': '±',
        r'&times;': '×',
        r'&divide;': '÷',
        r'&cent;': '¢',
        r'&pound;': '£',
        r'&euro;': '€',
        r'&yen;': '¥',
        r'&part;': '∂',
        r'&sum;': '∑',
        r'&prod;': '∏',
        r'&pi;': 'π',
        r'&theta;': 'θ',
        r'&phi;': 'φ',
        r'&omega;': 'ω',
        r'&alpha;': 'α',
        r'&beta;': 'β',
        r'&gamma;': 'γ',
        r'&delta;': 'δ',
        r'&epsilon;': 'ε',
        r'&sigma;': 'σ',
        r'&tau;': 'τ',
        r'&mu;': 'μ',
        r'&lambda;': 'λ',
        r'&upsilon;': 'υ',
        r'&phi;': 'φ',
        r'&chi;': 'χ',
        r'&psi;': 'ψ',
        r'&omega;': 'ω',
        r'&Alpha;': 'Α',
        r'&Beta;': 'Β',
        r'&Gamma;': 'Γ',
        r'&Delta;': 'Δ',
        r'&Epsilon;': 'Ε',
        r'&Zeta;': 'Ζ',
        r'&Eta;': 'Η',
        r'&Theta;': 'Θ',
        r'&Iota;': 'Ι',
        r'&Kappa;': 'Κ',
        r'&Lambda;': 'Λ',
        r'&Mu;': 'Μ',
        r'&Nu;': 'Ν',
        r'&Xi;': 'Ξ',
        r'&Omicron;': 'Ο',
        r'&Pi;': 'Π',
        r'&Rho;': 'Ρ',
        r'&Sigma;': 'Σ',
        r'&Tau;': 'Τ',
        r'&Upsilon;': 'Υ',
        r'&Phi;': 'Φ',
        r'&Chi;': 'Χ',
        r'&Psi;': 'Ψ',
        r'&Omega;': 'Ω',
        r'&le;': '≤',
        r'&ge;': '≥',
        r'&ne;': '≠',
        r'&equiv;': '≡',
        r'&approx;': '≈',
        r'&sim;': '~',
        r'&perp;': '⊥',
        r'&parallel;': '∥',
        r'&deg;': '°',
        r'&prime;': '′',
        r'&Prime;': '″',
        r'&micro;': 'μ',
        r'&ohm;': 'Ω',
        r'&ang;': '∠',
        r'&nabla;': '∇',
        r'&infin;': '∞',
        r'&partial;': '∂',
        r'&int;': '∫',
        r'&sum;': '∑',
        r'&prod;': '∏',
        r'&sqrt;': '√',
        r'&cube;': '∛',
        r'&frac12;': '½',
        r'&frac14;': '¼',
        r'&frac34;': '¾',
        r'&frac13;': '⅓',
        r'&frac23;': '⅔',
        r'&frac15;': '⅕',
        r'&frac25;': '⅖',
        r'&frac35;': '⅗',
        r'&frac45;': '⅘',
        r'&frac16;': '⅙',
        r'&frac56;': '⅚',
        r'&frac18;': '⅛',
        r'&frac38;': '⅜',
        r'&frac58;': '⅝',
        r'&frac78;': '⅞',
        r'&laquo;': '«',
        r'&raquo;': '»',
        r'&dagger;': '†',
        r'&Dagger;': '‡',
        r'&hellip;': '…',
        r'&lsaquo;': '‹',
        r'&rsaquo;': '›',
        r'&sbquo;': '‚',
        r'&bdquo;': '„',
        r'&lrm;': '',
        r'&rlm;': '',
        r'&zwj;': '',
        r'&zwnj;': '',
    }
    
    for entity, replacement in html_entities.items():
        content = re.sub(entity, replacement, content)
    
    # 移除HTML注释
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)
    
    # 处理pre标签（保留代码块格式）
    def pre_tag_replace(match):
        pre_content = match.group(1)
        return f'\n```\n{pre_content}\n```\n'
    content = re.sub(r'\s*<pre[^>]*>\s*<code[^>]*>(.*?)</code>\s*</pre>\s*', pre_tag_replace, content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'\s*<pre[^>]*>\s*(.*?)\s*</pre>\s*', pre_tag_replace, content, flags=re.IGNORECASE | re.DOTALL)
    
    # 处理blockquote标签（转换为Markdown引用）
    def blockquote_replace(match):
        quote_content = match.group(1)
        lines = quote_content.split('\n')
        quoted_lines = ['> ' + line if line.strip() else '>' for line in lines]
        return '\n' + '\n'.join(quoted_lines) + '\n'
    content = re.sub(r'\s*<blockquote[^>]*>\s*(.*?)\s*</blockquote>\s*', blockquote_replace, content, flags=re.IGNORECASE | re.DOTALL)
    
    # 处理code标签（保留代码格式）
    content = re.sub(r'\s*<code[^>]*>\s*(.*?)\s*</code>\s*', r'`\1`', content, flags=re.IGNORECASE | re.DOTALL)
    
    # 先转换主要的格式化标签为Markdown，这样它们不会被后续的标签移除
    
    # 转换标题标签为Markdown标题
    for i in range(1, 7):
        tag = f'h{i}'
        # 捕获所有标题内容，包括换行符
        content = re.sub(rf'<{tag}[^>]*>(.*?)</{tag}>', rf'{"#"*i} \1\n\n', content, flags=re.IGNORECASE | re.DOTALL)
        # 处理异常情况：不匹配的结束标签
        content = re.sub(rf'<{tag}[^>]*>(.*?)</[^>]+>', rf'{"#"*i} \1\n\n', content, flags=re.IGNORECASE | re.DOTALL)
        # 处理只有开始标签没有结束标签的情况
        content = re.sub(rf'<{tag}[^>]*>(.*?)(?:$|<)', rf'{"#"*i} \1\n\n', content, flags=re.IGNORECASE | re.DOTALL)

    # 然后处理段落标签转换为换行
    content = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', content, flags=re.IGNORECASE | re.DOTALL)
    # 确保标题和段落之间有正确的换行符
    content = re.sub(r'(### .*?)\n(.*?)', r'\1\n\n\2', content, flags=re.DOTALL)
    
    # 列表标签转换为Markdown列表
    # 确保ul/ol标签前后都有换行符，特别是当它们紧跟在其他内容后面时
    content = re.sub(r'(?<!\n)\s*<ul[^>]*>\s*', r'\n\n', content, flags=re.IGNORECASE)
    content = re.sub(r'\s*</ul>\s*(?!\n)', r'\n\n', content, flags=re.IGNORECASE)
    content = re.sub(r'(?<!\n)\s*<ol[^>]*>\s*', r'\n\n', content, flags=re.IGNORECASE)
    content = re.sub(r'\s*</ol>\s*(?!\n)', r'\n\n', content, flags=re.IGNORECASE)
    content = re.sub(r'\s*<li[^>]*>\s*(.*?)\s*</li>\s*', r'- \1\n', content, flags=re.IGNORECASE | re.DOTALL)
    
    # 特殊处理段落结束后直接跟列表的情况
    content = re.sub(r'(\.)\s*(- )', r'\1\n\2', content)
    # 更通用的处理：确保在任何文本后直接跟列表项时都添加换行符
    content = re.sub(r'([^\n])\s*(- )', r'\1\n\2', content)
    
    # 处理加粗和斜体标签
    content = re.sub(r'\s*<(strong|b)[^>]*>\s*(.*?)\s*</(strong|b)>\s*', r'**\2**', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'\s*<(em|i)[^>]*>\s*(.*?)\s*</(em|i)>\s*', r'*\2*', content, flags=re.IGNORECASE | re.DOTALL)
    
    # 处理链接标签（简单保留文本内容）
    content = re.sub(r'\s*<a[^>]*>\s*(.*?)\s*</a>\s*', r'\1', content, flags=re.IGNORECASE | re.DOTALL)
    
    # 处理图片标签（移除）
    content = re.sub(r'\s*<img[^>]*>\s*', '', content, flags=re.IGNORECASE)
    
    # 处理<br>标签为换行
    content = re.sub(r'\s*<br[^>]*>\s*', r'\n', content, flags=re.IGNORECASE)
    
    # 处理<hr>标签为分隔线
    content = re.sub(r'\s*<hr[^>]*>\s*', r'\n---\n', content, flags=re.IGNORECASE)
    
    # 处理div和span标签
    content = re.sub(r'\s*<div[^>]*>\s*', r'\n\n', content, flags=re.IGNORECASE)
    content = re.sub(r'\s*</div>\s*', r'\n\n', content, flags=re.IGNORECASE)
    content = re.sub(r'\s*<span[^>]*>\s*', '', content, flags=re.IGNORECASE)
    content = re.sub(r'\s*</span>\s*', '', content, flags=re.IGNORECASE)
    
    # 处理表格标签（简单转换为文本）
    content = re.sub(r'\s*<table[^>]*>\s*', r'\n', content, flags=re.IGNORECASE)
    content = re.sub(r'\s*</table>\s*', r'\n', content, flags=re.IGNORECASE)
    content = re.sub(r'\s*<tr[^>]*>\s*', r'\n', content, flags=re.IGNORECASE)
    content = re.sub(r'\s*</tr>\s*', r'\n', content, flags=re.IGNORECASE)
    content = re.sub(r'\s*<td[^>]*>\s*(.*?)\s*</td>\s*', r'| \1 ', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'\s*<th[^>]*>\s*(.*?)\s*</th>\s*', r'| **\1** ', content, flags=re.IGNORECASE | re.DOTALL)
    
    # 处理其他块级元素
    block_elements = ['section', 'article', 'header', 'footer', 'nav', 'aside', 'main', 'figure', 'figcaption', 'form', 'fieldset', 'legend', 'label', 'button', 'select', 'option', 'optgroup', 'datalist', 'textarea', 'input', 'progress', 'meter', 'details', 'summary', 'dialog']
    for element in block_elements:
        content = re.sub(rf'\s*<{element}[^>]*>\s*', r'\n\n', content, flags=re.IGNORECASE)
        content = re.sub(rf'\s*</{element}>\s*', r'\n\n', content, flags=re.IGNORECASE)
    
    # 处理内联元素
    inline_elements = ['u', 's', 'sup', 'sub', 'kbd', 'var', 'samp', 'cite', 'dfn', 'abbr', 'mark', 'small', 'del', 'ins', 'ruby', 'rt', 'rp', 'bdi', 'bdo', 'data', 'time', 'output', 'link', 'meta', 'base', 'style', 'script', 'noscript', 'template', 'slot', 'canvas', 'svg', 'math']
    for element in inline_elements:
        content = re.sub(rf'\s*<{element}[^>]*>\s*', '', content, flags=re.IGNORECASE)
        content = re.sub(rf'\s*</{element}>\s*', '', content, flags=re.IGNORECASE)
    
    # 处理自闭合标签
    content = re.sub(r'\s*<[^>]+/>\s*', '', content, flags=re.IGNORECASE)
    
    # 处理可能的未闭合标签
    content = re.sub(r'<[^>]*$', '', content)  # 处理行尾的未闭合标签
    content = re.sub(r'<[^>]+\s*$', '', content)  # 处理文件末尾的未闭合标签
    
    # 多次去除所有剩余的HTML标签，确保彻底清除嵌套标签
    for _ in range(10):  # 增加到10次处理，确保嵌套标签被彻底清除
        content = re.sub(r'<[^>]+>', '', content, flags=re.DOTALL)
    
    # 清理可能的实体残留
    content = re.sub(r'&[^;]+;', '', content)
    
    # 清理多余的空行
    content = re.sub(r'\n{3,}', r'\n\n', content)
    
    # 清理行首行尾的空格
    content = re.sub(r'^\s+', '', content, flags=re.MULTILINE)
    content = re.sub(r'\s+$', '', content, flags=re.MULTILINE)
    
    return content.strip()


@router.post("/upload", response_model=FileUploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """上传文档文件并提取文本内容"""
    try:
        # 检查文件类型
        allowed_types = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
        
        if file.content_type not in allowed_types:
            return FileUploadResponse(
                success=False,
                message="不支持的文件类型，请上传PDF或Word文档"
            )
        
        # 处理文件并提取文本
        file_content = await FileService.process_uploaded_file(file)
        
        return FileUploadResponse(
            success=True,
            message=f"文件 {file.filename} 上传成功",
            file_content=file_content
        )
        
    except Exception as e:
        return FileUploadResponse(
            success=False,
            message=f"文件处理失败: {str(e)}"
        )


@router.post("/analyze-stream")
async def analyze_document_stream(request: AnalysisRequest):
    """流式分析文档内容"""
    try:
        # 加载配置
        config = config_manager.load_config()
        
        if not config.get('api_key'):
            raise HTTPException(status_code=400, detail="请先配置OpenAI API密钥")

        # 创建OpenAI服务实例
        openai_service = OpenAIService()
        
        async def generate():
            # 构建分析提示词
            if request.analysis_type == AnalysisType.OVERVIEW:
                system_prompt = """你是一个专业的标书撰写专家。请分析用户发来的招标文件，提取并总结项目概述信息。
            
请重点关注以下方面：
1. 项目名称和基本信息
2. 项目背景和目的
3. 项目规模和预算
4. 项目时间安排
5. 项目要实施的具体内容
6. 主要技术特点
7. 其他关键要求

工作要求：
1. 保持提取信息的全面性和准确性，尽量使用原文内容，不要自己编写
2. 只关注与项目实施有关的内容，不提取商务信息
3. 直接返回整理好的项目概述，除此之外不返回任何其他内容
"""
            else:  # requirements
                system_prompt = """你是一名专业的招标文件分析师，擅长从复杂的招标文档中高效提取“技术评分项”相关内容。请严格按照以下步骤和规则执行任务：
### 1. 目标定位
- 重点识别文档中与“技术评分”、“评标方法”、“评分标准”、“技术参数”、“技术要求”、“技术方案”、“技术部分”或“评审要素”相关的章节（如“第X章 评标方法”或“附件X：技术评分表”）。
- 一定不要提取商务、价格、资质等于技术类评分项无关的条目。
### 2. 提取内容要求
对每一项技术评分项，按以下结构化格式输出（若信息缺失，标注“未提及”），如果评分项不够明确，你需要根据上下文分析并也整理成如下格式：
【评分项名称】：<原文描述，保留专业术语>
【权重/分值】：<具体分值或占比，如“30分”或“40%”>
【评分标准】：<详细规则，如“≥95%得满分，每低1%扣0.5分”>
【数据来源】：<文档中的位置，如“第5.2.3条”或“附件3-表2”>

### 3. 处理规则
- **模糊表述**：有些招标文件格式不是很标准，没有明确的“技术评分表”，但一定都会有“技术评分”相关内容，请根据上下文判断评分项。
- **表格处理**：若评分项以表格形式呈现，按行提取，并标注“[表格数据]”。
- **分层结构**：若存在二级评分项（如“技术方案→子项1、子项2”），用缩进或编号体现层级关系。
- **单位统一**：将所有分值统一为“分”或“%”，并注明原文单位（如原文为“20点”则标注“[原文：20点]”）。

### 4. 输出示例
【评分项名称】：系统可用性 
【权重/分值】：25分 
【评分标准】：年平均故障时间≤1小时得满分；每增加1小时扣2分，最高扣10分。 
【数据来源】：附件4-技术评分细则（第3页） 

【评分项名称】：响应时间
【权重/分分】：15分 [原文：15%]
【评分标准】：≤50ms得满分；每增加10ms扣1分。
【数据来源】：第6.1.2条

### 5. 验证步骤
提取完成后，执行以下自检：
- [ ] 所有技术评分项是否覆盖（无遗漏）？
- [ ] 是否错误提取商务、价格、资质等于技术类评分项无关的条目？
- [ ] 权重总和是否与文档声明的技术分总分一致（如“技术部分共60分”）？

直接返回提取结果，除此之外不输出任何其他内容
"""
            
            analysis_type_cn = "项目概述" if request.analysis_type == AnalysisType.OVERVIEW else "技术评分要求"
            user_prompt = f"请分析以下招标文件内容，提取{analysis_type_cn}信息：\n\n{request.file_content}"
            
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
            
            # 流式返回分析结果
            async for chunk in openai_service.stream_chat_completion(messages, temperature=0.3):
                yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
            
            # 发送结束信号
            yield "data: [DONE]\n\n"
        
        return sse_response(generate())
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档分析失败: {str(e)}")


@router.post("/export-word")
async def export_word(request: WordExportRequest):
    """根据目录数据导出Word文档"""
    try:
        doc = docx.Document()

        # 统一设置文档的基础字体为宋体，取消普通段落默认加粗
        try:
            styles = doc.styles
            base_styles = ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]
            for style_name in base_styles:
                if style_name in styles:
                    style = styles[style_name]
                    font = style.font
                    font.name = "宋体"
                    # 设置中文字体
                    if style._element.rPr is None:
                        style._element._add_rPr()
                    rpr = style._element.rPr
                    rpr.rFonts.set(qn("w:eastAsia"), "宋体")
                    if style_name == "Normal":
                        font.bold = False
        except Exception:
            # 字体设置失败不影响文档生成，忽略
            pass

        # AI 生成声明
        p = doc.add_paragraph()
        run = p.add_run("内容由AI生成")
        run.italic = True
        run.font.size = Pt(9)
        set_run_font_simsun(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 文档标题
        title = request.project_name or "投标技术文件"
        # 清理标题中的HTML标签
        title = clean_html_tags(title)
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        set_run_font_simsun(title_run)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 项目概述
        if request.project_overview:
            heading = doc.add_heading("项目概述", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_font_simsun(heading)
            # 清理项目概述中的HTML标签
            cleaned_overview = clean_html_tags(request.project_overview)
            overview_p = doc.add_paragraph(cleaned_overview)
            set_paragraph_font_simsun(overview_p)
            overview_p_format = overview_p.paragraph_format
            overview_p_format.space_after = Pt(12)

        # 简单的 Markdown 段落解析：支持标题、列表、表格和基础加粗/斜体
        def add_markdown_runs(para: docx.text.paragraph.Paragraph, text: str) -> None:
            """在指定段落中追加 markdown 文本的 runs"""
            pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`)"
            parts = re.split(pattern, text)
            for part in parts:
                if not part:
                    continue
                run = para.add_run()
                # 加粗
                if part.startswith("**") and part.endswith("**") and len(part) > 4:
                    run.text = part[2:-2]
                    run.bold = True
                # 斜体
                elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                    run.text = part[1:-1]
                    run.italic = True
                # 行内代码：这里只去掉反引号
                elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                    run.text = part[1:-1]
                else:
                    run.text = part
                # 确保字体为宋体
                set_run_font_simsun(run)

        def add_markdown_paragraph(text: str) -> None:
            """将一段 Markdown 文本解析为一个普通段落，保留加粗/斜体效果"""
            para = doc.add_paragraph()
            add_markdown_runs(para, text)
            para.paragraph_format.space_after = Pt(6)

        def parse_markdown_blocks(content: str):
            """
            识别 Markdown 内容中的块级元素，返回结构化的 block 列表：
            - ('list', items)        items: [(kind, num_str, text), ...]
            - ('table', rows)        rows: [text, ...]
            - ('heading', level, text)
            - ('paragraph', text)
            """
            blocks = []
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i].rstrip("\r").strip()
                if not line:
                    i += 1
                    continue

                # 列表项（有序/无序）
                if line.startswith("- ") or line.startswith("* ") or re.match(r"^\d+\.\s", line):
                    # items: (kind, number, text)
                    items = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        # 无序列表
                        if stripped.startswith("- ") or stripped.startswith("* "):
                            text = re.sub(r"^[-*]\s+", "", stripped).strip()
                            if text:
                                items.append(("unordered", None, text))
                            i += 1
                            continue
                        # 有序列表（1. xxx）
                        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
                        if m_num:
                            num_str, text = m_num.groups()
                            text = text.strip()
                            if text:
                                items.append(("ordered", num_str, text))
                            i += 1
                            continue
                        break

                    if items:
                        blocks.append(("list", items))
                    continue

                # 表格（简化为每行一个段落，单元格用 | 分隔）
                if "|" in line:
                    rows = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        if "|" in stripped:
                            # 跳过仅由 - 和 | 组成的分隔行
                            if not re.match(r"^\|?[-\s\|]+\|?$", stripped):
                                cells = [c.strip() for c in stripped.split("|")]
                                row_text = " | ".join([c for c in cells if c])
                                if row_text:
                                    rows.append(row_text)
                            i += 1
                        else:
                            break
                    if rows:
                        blocks.append(("table", rows))
                    continue

                # Markdown 标题（# / ## / ###）
                if line.startswith("#"):
                    m = re.match(r"^(\#+)\s*(.*)$", line)
                    if m:
                        level_marks, title_text = m.groups()
                        level = min(len(level_marks), 3)
                        blocks.append(("heading", level, title_text.strip()))
                    i += 1
                    continue

                # 普通段落：合并连续的普通行
                para_lines = []
                while i < len(lines):
                    raw = lines[i].rstrip("\r")
                    stripped = raw.strip()
                    if (
                        stripped
                        and not stripped.startswith("-")
                        and not stripped.startswith("*")
                        and "|" not in stripped
                        and not stripped.startswith("#")
                    ):
                        para_lines.append(stripped)
                        i += 1
                    else:
                        break
                if para_lines:
                    text = " ".join(para_lines)
                    blocks.append(("paragraph", text))
                else:
                    i += 1

            return blocks

        def render_markdown_blocks(blocks) -> None:
            """将结构化的 Markdown blocks 渲染到文档"""
            for block in blocks:
                kind = block[0]
                if kind == "list":
                    items = block[1]
                    for item_kind, num_str, text in items:
                        p = doc.add_paragraph()
                        if item_kind == "unordered":
                            # 使用“• ”模拟项目符号
                            run = p.add_run("• ")
                            set_run_font_simsun(run)
                        else:
                            # 有序列表：输出 "1. " 这样的前缀
                            prefix = f"{num_str}."
                            run = p.add_run(prefix + " ")
                            set_run_font_simsun(run)
                        # 紧跟在同一段落中追加列表文本
                        add_markdown_runs(p, text)
                elif kind == "table":
                    rows = block[1]
                    for row in rows:
                        add_markdown_paragraph(row)
                elif kind == "heading":
                    _, level, text = block
                    heading = doc.add_heading(text, level=level)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_paragraph_font_simsun(heading)
                elif kind == "paragraph":
                    _, text = block
                    add_markdown_paragraph(text)

        def add_markdown_content(content: str) -> None:
            """解析并渲染 Markdown 文本到文档"""
            # 清理HTML标签，转换为纯Markdown内容
            content = clean_html_tags(content)
            blocks = parse_markdown_blocks(content)
            render_markdown_blocks(blocks)

        # 应用样式配置到文档
        def apply_style_config():
            """应用前端传递的样式配置"""
            if not request.styleConfig:
                return

            # 获取标题格式设置
            title_format = request.styleConfig.get('titleFormat', '第一章/第一节')
            chapter_start = request.styleConfig.get('chapterStart', '第一章')

            # 根据标题格式调整章节编号
            # 这里可以根据需要扩展更多的编号格式处理
            return {
                'title_format': title_format,
                'chapter_start': chapter_start
            }

        # 获取样式配置
        style_config = apply_style_config()

        # 递归构建文档内容（章节和内容）
        def add_outline_items(items, level: int = 1):
            for item in items:
                # 章节标题
                if level <= 3:
                    # 清理标题中的HTML标签
                    clean_title = clean_html_tags(item.title)
                    heading = doc.add_heading(f"{item.id} {clean_title}", level=level)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for hr in heading.runs:
                        hr.font.name = "宋体"
                        rr = hr._element.rPr
                        if rr is not None and rr.rFonts is not None:
                            rr.rFonts.set(qn("w:eastAsia"), "宋体")
                    
                    # 应用样式配置中的标题格式
                    if request.styleConfig:
                        level_key = f'level{level}'
                        if level_key in request.styleConfig:
                            level_style = request.styleConfig[level_key]
                            # 设置对齐方式
                            align_map = {
                                'left': WD_ALIGN_PARAGRAPH.LEFT,
                                'center': WD_ALIGN_PARAGRAPH.CENTER,
                                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                            }
                            if level_style.get('align') in align_map:
                                heading.alignment = align_map[level_style['align']]
                            
                            # 设置字号
                            if level_style.get('size'):
                                size_map = {
                                    '二号': 22,
                                    '三号': 16,
                                    '四号': 14,
                                    '五号': 12,
                                    '小四号': 12,
                                    '小五号': 9
                                }
                                if level_style['size'] in size_map:
                                    for hr in heading.runs:
                                        hr.font.size = Pt(size_map[level_style['size']])
                            
                            # 设置加粗
                            if 'bold' in level_style:
                                for hr in heading.runs:
                                    hr.font.bold = level_style['bold']
                else:
                    # 清理标题中的HTML标签
                    clean_title = clean_html_tags(item.title)
                    para = doc.add_paragraph()
                    run = para.add_run(f"{item.id} {clean_title}")
                    run.bold = True
                    run.font.name = "宋体"
                    rr = run._element.rPr
                    if rr is not None and rr.rFonts is not None:
                        rr.rFonts.set(qn("w:eastAsia"), "宋体")
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(3)

                # 叶子节点内容
                if not item.children:
                    content = item.content or ""
                    if content.strip():
                        add_markdown_content(content)
                else:
                    add_outline_items(item.children, level + 1)

        add_outline_items(request.outline)

        # 输出到内存并返回
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{request.project_name or '标书文档'}.docx"
        # 使用 RFC 5987 格式对文件名进行 URL 编码，避免非 ASCII 字符导致的编码错误
        encoded_filename = quote(filename)
        content_disposition = f"attachment; filename*=UTF-8''{encoded_filename}"
        headers = {
            "Content-Disposition": content_disposition
        }

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except Exception as e:
        # 打印详细错误信息到控制台，方便排查
        import traceback
        print("导出Word失败:", str(e))
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"导出Word失败: {str(e)}")


# 测试函数：验证clean_html_tags函数的修复效果
def test_clean_html_tags():
    """测试HTML标签清理功能"""
    import re
    
    # 直接导入clean_html_tags函数
    from app.routers.document import clean_html_tags
    clean_html_tags_func = clean_html_tags
    
    # 测试用例
    test_cases = [
        # 基本HTML标签
        ("<h3>测试标题</h3>", "### 测试标题"),
        ("<p>这是一个段落</p>", "这是一个段落"),
        ("<strong>加粗文本</strong>", "**加粗文本**"),
        ("<em>斜体文本</em>", "*斜体文本*"),
        
        # 嵌套HTML标签
        ("<p>这是一个<strong>加粗</strong>的段落</p>", "这是一个**加粗**的段落"),
        ("<div><h3>标题</h3><p>内容</p></div>", "### 标题\n内容"),
        ("<p>这是<em>斜体<strong>加粗斜体</strong></em>文本</p>", "这是*斜体**加粗斜体***文本"),
        
        # HTML实体
        ("<p>测试&amp;测试</p>", "测试&测试"),
        ("<p>测试&nbsp;空格</p>", "测试 空格"),
        ("<p>© 版权所有</p>", "© 版权所有"),
        
        # 特殊HTML标签
        ("<pre><code>print('hello world')</code></pre>", "```\nprint('hello world')\n```"),
        ("<blockquote>引用文本</blockquote>", "> 引用文本"),
        ("<code>code</code>", "`code`"),
        
        # 不规范HTML
        ("<h3>测试标题</h3><p>段落</p>", "### 测试标题\n段落"),
        ("<h3>测试标题<p>段落</h3>", "### 测试标题段落"),
        ("<h3>测试标题</p>", "### 测试标题"),
        
        # 多标签混合
        ("<div><h3>项目概述</h3><p>这是一个<em>详细</em>的项目概述，包含<strong>重要</strong>信息。</p><ul><li>点1</li><li>点2</li></ul></div>", "### 项目概述\n这是一个*详细*的项目概述，包含**重要**信息。\n- 点1\n- 点2"),
        
        # 用户反馈中提到的问题：h3标签残留
        ("<h3>技术方案</h3><p>详细内容</p>", "### 技术方案\n详细内容"),
    ]
    
    # 运行测试
    passed = 0
    failed = 0
    for i, (input_html, expected) in enumerate(test_cases, 1):
        result = clean_html_tags(input_html)
        if result == expected:
            print(f"测试 {i} 通过:")
            print(f"  输入: {input_html[:50]}{'...' if len(input_html) > 50 else ''}")
            print(f"  输出: {result[:50]}{'...' if len(result) > 50 else ''}")
            passed += 1
        else:
            print(f"测试 {i} 失败:")
            print(f"  输入: {input_html}")
            print(f"  期望: {expected}")
            print(f"  实际: {result}")
            failed += 1
        print()
    
    print(f"测试完成: {passed} 通过, {failed} 失败")
    return passed, failed


# 如果直接运行此文件，执行测试
if __name__ == "__main__":
    test_clean_html_tags()
